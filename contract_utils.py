from __future__ import annotations

import csv
import io
import re
import shutil
import subprocess
import tempfile
from copy import deepcopy
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import BinaryIO

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


DEFAULT_TEMPLATE_PATH = Path(__file__).parent / "templates" / "LTO-Template-Formatted-v2.docx"
DEFAULT_PROPERTY_CSV_PATH = Path(__file__).parent / "templates" / "Apartments.csv"
FIXED_COMPANY_NAME = "VJProp LTD"
FIXED_OWNER_REGISTERED_ADDRESS = "2 - 4 Ireton Street, BT7 1LH, Belfast."
FIXED_AUTHORISED_SIGNATORY_NAME = "Colin Johnston"
FIXED_OWNER_MOBILE_NUMBER = ""

PLACEHOLDER_DESCRIPTIONS = {
    "Agreement Date": "The date on which this agreement is executed.",
    "Full Name of Primary Licensee": "Full legal name of the first tenant.",
    "Date of Birth": "Date of birth for the primary tenant.",
    "National Insurance Number": "National Insurance number for the primary tenant.",
    "Full Name of Secondary Licensee / Guarantor": "Full legal name of the second tenant or guarantor.",
    "Company Name": "Registered name of the landlord or owning company.",
    "Owner's Registered Address": "Registered address of the landlord or owning company.",
    "Authorised Signatory Name": "Name of the person signing on behalf of the landlord.",
    "Full Property Address": "Property address used wherever the contract references the apartment.",
    "Development Name": "Development or building name.",
    "Deposit Amount in £": "Security deposit in pounds sterling.",
    "Monthly Licence Fee in £": "Monthly rent in pounds sterling.",
    "First Payment Date": "Start date of the tenancy and first payment date.",
    "Payment Day": "Recurring payment day of the month, shown as a British ordinal.",
    "End Date of Term": "Fixed term end date.",
    "Notice Period": "Written notice period for early termination.",
    "Minimum Occupation Period Before Notice": "Minimum occupation period before notice can be served.",
    "Owner's Mobile Number": "Landlord contact number for notices.",
}

TENANT_DOC_TYPES = (
    "",
    "UK/Irish Passport",
    "BRP",
    "UK/Irish Driving Licence",
    "eVisa",
    "Oversea Passport",
)

PAYMENT_DAY_OPTIONS = [f"{value:02d}" for value in range(1, 32)]


@dataclass(slots=True)
class PropertyRecord:
    title: str
    building_name: str
    apartment_number: str
    full_address: str
    bedrooms: int | None
    bathrooms: int | None

    @property
    def apartment_description(self) -> str:
        bedroom_part = pluralise_count(self.bedrooms, "Bedroom")
        bathroom_part = pluralise_count(self.bathrooms, "Bathroom")
        parts = [part for part in (bedroom_part, bathroom_part, "Kitchen/Living Area") if part]
        return ", ".join(parts)


@dataclass(slots=True)
class TenantData:
    full_name: str
    date_of_birth: str
    national_insurance_number: str
    id_document_type: str
    id_number: str


def pluralise_count(value: int | None, noun: str) -> str:
    if value is None:
        return ""
    if value == 1:
        return f"One {noun}"
    if value == 2:
        return f"Two {noun}s"
    if value == 3:
        return f"Three {noun}s"
    return f"{value} {noun}s"


def read_properties_from_bytes(content: bytes) -> list[PropertyRecord]:
    handle = io.StringIO(content.decode("utf-8-sig"))
    reader = csv.DictReader(handle)
    properties = []
    for row in reader:
        if not row.get("Title"):
            continue
        properties.append(
            PropertyRecord(
                title=row.get("Title", "").strip(),
                building_name=row.get("BuildingName", "").strip(),
                apartment_number=row.get("ApartmentNumber", "").strip(),
                full_address=row.get("FullAddress", "").strip(),
                bedrooms=parse_int(row.get("Bedrooms", "")),
                bathrooms=parse_int(row.get("Bathrooms", "")),
            )
        )
    return properties


def read_properties(csv_path: Path) -> list[PropertyRecord]:
    with csv_path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        properties = []
        for row in reader:
            if not row.get("Title"):
                continue
            properties.append(
                PropertyRecord(
                    title=row.get("Title", "").strip(),
                    building_name=row.get("BuildingName", "").strip(),
                    apartment_number=row.get("ApartmentNumber", "").strip(),
                    full_address=row.get("FullAddress", "").strip(),
                    bedrooms=parse_int(row.get("Bedrooms", "")),
                    bathrooms=parse_int(row.get("Bathrooms", "")),
                )
            )
    return properties


def parse_int(raw_value: str) -> int | None:
    value = (raw_value or "").strip()
    if not value:
        return None
    try:
        return int(value)
    except ValueError:
        return None


def ordinalise_day(day_value: str | int) -> str:
    day = int(day_value)
    if 10 <= day % 100 <= 20:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    return f"{day}{suffix}"


def format_contract_date(value: date) -> str:
    return value.strftime("%d/%m/%Y")


def format_slash_date(value: date) -> str:
    return value.strftime("%d/%m/%Y")


def discover_template_placeholders_from_bytes(content: bytes) -> list[str]:
    document = Document(io.BytesIO(content))
    placeholders: set[str] = set()
    for paragraph in iter_all_paragraphs(document):
        placeholders.update(re.findall(r"\[([^\]]+)\]", paragraph.text))
    return sorted(clean_placeholder_name(name) for name in placeholders if should_expose_placeholder(name))


def discover_template_placeholders(template_path: Path) -> list[str]:
    document = Document(template_path)
    placeholders: set[str] = set()
    for paragraph in iter_all_paragraphs(document):
        placeholders.update(re.findall(r"\[([^\]]+)\]", paragraph.text))
    return sorted(clean_placeholder_name(name) for name in placeholders if should_expose_placeholder(name))


def should_expose_placeholder(raw_name: str) -> bool:
    return ", e.g." not in raw_name and "including Postcode" not in raw_name


def clean_placeholder_name(raw_name: str) -> str:
    name = raw_name.strip()
    if ", e.g." in name:
        name = name.split(", e.g.", 1)[0].strip()
    if " including Postcode" in name:
        name = name.replace(" including Postcode", "").strip()
    return name


def build_placeholder_mapping(
    *,
    agreement_date: date,
    start_date: date,
    end_date: date,
    payment_day: str,
    deposit_amount: str,
    monthly_rent: str,
    notice_period: str,
    minimum_occupation_period: str,
    company_name: str,
    owner_registered_address: str,
    owner_mobile_number: str,
    authorised_signatory_name: str,
    property_record: PropertyRecord,
    tenants: list[TenantData],
) -> dict[str, str]:
    first_tenant = tenants[0]
    second_tenant = tenants[1] if len(tenants) > 1 else TenantData("", "", "", TENANT_DOC_TYPES[0], "")
    property_address = property_record.full_address
    mapping = {
        "Agreement Date": format_slash_date(agreement_date),
        "Full Name of Primary Licensee": first_tenant.full_name,
        "Date of Birth": first_tenant.date_of_birth,
        "National Insurance Number": first_tenant.national_insurance_number,
        "Full Name of Secondary Licensee / Guarantor": second_tenant.full_name,
        "Company Name": company_name.strip(),
        "Owner's Registered Address": owner_registered_address.strip(),
        "Authorised Signatory Name": authorised_signatory_name.strip(),
        "Full Property Address": property_address,
        "Development Name": property_record.building_name,
        "Apartment Description": "",
        "Deposit Amount in £": normalise_currency(deposit_amount),
        "Monthly Licence Fee in £": normalise_currency(monthly_rent),
        "First Payment Date": format_contract_date(start_date),
        "Payment Day": ordinalise_day(payment_day),
        "End Date of Term": format_contract_date(end_date),
        "Notice Period": notice_period.strip(),
        "Minimum Occupation Period Before Notice": minimum_occupation_period.strip(),
        "Owner's Mobile Number": owner_mobile_number.strip(),
    }
    return mapping


def normalise_currency(raw_value: str) -> str:
    value = raw_value.strip().replace("£", "").replace(",", "")
    if not value:
        return ""
    try:
        number = float(value)
    except ValueError:
        return raw_value.strip()
    if number.is_integer():
        return f"£{int(number):,}"
    return f"£{number:,.2f}"


def render_contract_from_bytes(
    template_bytes: bytes,
    placeholder_mapping: dict[str, str],
    tenants: list[TenantData],
) -> bytes:
    document = Document(io.BytesIO(template_bytes))
    strip_data_input_list(document)
    strip_apartment_description_clause(document)
    rename_payment_date_clause(document)
    if not placeholder_mapping.get("Owner's Mobile Number", "").strip():
        strip_owner_mobile_notice(document)
    apply_placeholder_mapping(document, placeholder_mapping)
    rebuild_party_section(document, tenants, placeholder_mapping)
    rebuild_execution_section(document, tenants, placeholder_mapping)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def render_contract(
    template_path: Path,
    placeholder_mapping: dict[str, str],
    tenants: list[TenantData],
) -> bytes:
    document = Document(template_path)
    strip_data_input_list(document)
    strip_apartment_description_clause(document)
    rename_payment_date_clause(document)
    if not placeholder_mapping.get("Owner's Mobile Number", "").strip():
        strip_owner_mobile_notice(document)
    apply_placeholder_mapping(document, placeholder_mapping)
    rebuild_party_section(document, tenants, placeholder_mapping)
    rebuild_execution_section(document, tenants, placeholder_mapping)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def apply_placeholder_mapping(document: DocumentObject, placeholder_mapping: dict[str, str]) -> None:
    replacements = build_replacement_pairs(placeholder_mapping)
    for paragraph in iter_all_paragraphs(document):
        replace_in_paragraph(paragraph, replacements)


def build_replacement_pairs(placeholder_mapping: dict[str, str]) -> list[tuple[str, str]]:
    pairs: list[tuple[str, str]] = []
    for key, value in placeholder_mapping.items():
        text_value = value or ""
        pairs.append((f"[{key}]", text_value))
        if key == "Apartment Description":
            pairs.append((f"[{key}, e.g. Two Bedrooms, One Bathroom, Kitchen/Living Area]", text_value))
        elif key == "First Payment Date":
            pairs.append((f"[{key}, e.g. DD MMM YYYY]", text_value))
        elif key == "Payment Day":
            pairs.append((f"[{key}, e.g. 1st / 3rd / 15th]", text_value))
        elif key == "End Date of Term":
            pairs.append((f"[{key}, e.g. DD MMM YYYY]", text_value))
        elif key == "Notice Period":
            pairs.append((f"[{key}, e.g. one calendar month]", text_value))
        elif key == "Minimum Occupation Period Before Notice":
            pairs.append((f"[{key}, e.g. one calendar month]", text_value))
        elif key == "Full Property Address":
            pairs.append((f"[{key} including Postcode]", text_value))
    return pairs


def iter_all_paragraphs(document: DocumentObject):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        yield from iter_table_paragraphs(table)
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for paragraph in section.footer.paragraphs:
            yield paragraph


def iter_table_paragraphs(table: Table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested_table in cell.tables:
                yield from iter_table_paragraphs(nested_table)


def replace_in_paragraph(paragraph: Paragraph, replacements: list[tuple[str, str]]) -> None:
    combined = "".join(run.text for run in paragraph.runs)
    if not combined:
        return

    updated = combined
    for placeholder, replacement in replacements:
        updated = updated.replace(placeholder, replacement)

    if updated == combined:
        return

    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = updated


def strip_data_input_list(document: DocumentObject) -> None:
    body = document._element.body
    delete_mode = False
    for child in list(body):
        if child.tag == qn("w:p"):
            text = "".join(node.text or "" for node in child.iter(qn("w:t")))
            if text.strip() == "DATA INPUT LIST":
                delete_mode = True
        if delete_mode and child.tag in {qn("w:p"), qn("w:tbl")}:
            body.remove(child)


def rebuild_party_section(
    document: DocumentObject,
    tenants: list[TenantData],
    placeholder_mapping: dict[str, str],
) -> None:
    owner_paragraph = next(
        paragraph
        for paragraph in document.paragraphs
        if '("the Owner")' in paragraph.text
    )
    owner_element = owner_paragraph._element
    body = document._element.body

    tenant_templates = []
    for paragraph in document.paragraphs:
        if '("the Licensee");' in paragraph.text:
            tenant_templates.append(paragraph)
    if not tenant_templates:
        return

    template_element = tenant_templates[0]._element
    previous = template_element

    for paragraph in tenant_templates:
        body.remove(paragraph._element)

    for index, tenant in enumerate(tenants, start=1):
        new_element = deepcopy(template_element)
        owner_element.addprevious(new_element)
        new_paragraph = Paragraph(new_element, tenant_templates[0]._parent)
        new_paragraph.style = tenant_templates[0].style
        detail_parts = []
        if tenant.date_of_birth.strip():
            detail_parts.append(f"Date of Birth: {tenant.date_of_birth.strip()}")
        if tenant.national_insurance_number.strip():
            detail_parts.append(
                f"National Insurance Number: {tenant.national_insurance_number.strip()}"
            )
        if tenant.id_document_type.strip():
            detail_parts.append(f"ID Document: {tenant.id_document_type.strip()}")
        if tenant.id_number.strip():
            detail_parts.append(f"ID Number: {tenant.id_number.strip()}")
        details = f", {', '.join(detail_parts)}" if detail_parts else ""
        new_text = f'{index}.  {tenant.full_name}{details} ("the Licensee");'
        set_paragraph_text(new_paragraph, new_text)
        previous = new_element

    owner_address = placeholder_mapping["Owner's Registered Address"]
    owner_text = (
        f"{len(tenants) + 1}.  {placeholder_mapping['Company Name']} of "
        f"{owner_address} (\"the Owner\")"
    )
    set_paragraph_text(owner_paragraph, owner_text)


def rebuild_execution_section(
    document: DocumentObject,
    tenants: list[TenantData],
    placeholder_mapping: dict[str, str],
) -> None:
    body = document._element.body
    execution_heading = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "EXECUTION"
    )
    start_element = execution_heading._element
    elements_to_remove = []
    deleting = False
    for child in list(body):
        if child is start_element:
            deleting = True
        if deleting and child.tag in {qn("w:p"), qn("w:tbl")}:
            elements_to_remove.append(child)
    for child in elements_to_remove:
        body.remove(child)

    paragraphs = [
        ("EXECUTION", execution_heading.style),
        (
            "This document has been executed as a deed and is delivered and takes effect on the date stated at the beginning of it.",
            None,
        ),
    ]
    append_paragraph(document, paragraphs[0][0], paragraphs[0][1])
    append_paragraph(document, paragraphs[1][0], paragraphs[1][1])

    for tenant in tenants:
        add_signature_block(document, f"Signed by the Licensee: {tenant.full_name}", tenant.full_name)

    add_signature_block(
        document,
        "Signed by the Owner / Authorised Signatory:",
        f"{placeholder_mapping['Authorised Signatory Name']} on behalf of {placeholder_mapping['Company Name']}",
    )


def append_paragraph(document: DocumentObject, text: str, style=None) -> Paragraph:
    paragraph = document.add_paragraph()
    if style is not None:
        paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def add_signature_block(document: DocumentObject, heading: str, printed_name: str) -> None:
    append_paragraph(document, "")
    append_paragraph(document, heading)
    append_paragraph(document, "Signature:  ____________________________________________")
    append_paragraph(document, "Date:       ____________________________________________")
    append_paragraph(document, f"Name (Print):  {printed_name}")


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def strip_apartment_description_clause(document: DocumentObject) -> None:
    for paragraph in document.paragraphs:
        if '"Apartment" —' in paragraph.text and "[Apartment Description" in paragraph.text:
            set_paragraph_text(
                paragraph,
                '"Apartment" —  The apartment addressed as [Full Property Address], '
                'forming part of the development known as [Development Name].',
            )
            return


def rename_payment_date_clause(document: DocumentObject) -> None:
    for paragraph in document.paragraphs:
        stripped = paragraph.text.strip()
        if stripped.startswith('"Payment Date"') and "[First Payment Date" in paragraph.text:
            set_paragraph_text(
                paragraph,
                '"First Payment Date" —  [First Payment Date]',
            )
            continue
        if stripped.startswith('"Payment Date"') and "[Payment Day" in paragraph.text:
            set_paragraph_text(
                paragraph,
                '"Payment Day" —  The [Payment Day, e.g. 1st / 3rd / 15th] of each month until termination of this agreement.',
            )


def strip_owner_mobile_notice(document: DocumentObject) -> None:
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith("13.2") and "[Owner's Mobile Number]" in paragraph.text:
            set_paragraph_text(
                paragraph,
                "13.2  A notice given under this agreement shall be validly served if sent by email.",
            )
            return


def extract_id_details(
    uploaded_file: BinaryIO,
    filename: str,
    document_type: str,
) -> dict[str, str]:
    if not shutil.which("tesseract"):
        raise RuntimeError("Tesseract is not installed on this machine.")

    suffix = Path(filename).suffix or ".png"
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as handle:
        temp_path = Path(handle.name)
        payload = uploaded_file.read()
        handle.write(payload)

    try:
        result = subprocess.run(
            ["tesseract", str(temp_path), "stdout"],
            capture_output=True,
            check=False,
            text=True,
        )
        if result.returncode != 0:
            raise RuntimeError(result.stderr.strip() or "OCR extraction failed.")
        text = normalise_ocr_text(result.stdout)
        return parse_id_text(text, document_type)
    finally:
        temp_path.unlink(missing_ok=True)
        if hasattr(uploaded_file, "seek"):
            uploaded_file.seek(0)


def normalise_ocr_text(text: str) -> str:
    text = text.replace("\u2014", "-")
    text = re.sub(r"[ \t]+", " ", text)
    return text


def parse_id_text(text: str, document_type: str) -> dict[str, str]:
    if document_type in {"UK/Irish Passport", "Oversea Passport"}:
        return parse_uk_passport(text)
    if document_type == "BRP":
        return parse_brp(text)
    if document_type == "eVisa":
        return parse_evisa(text)
    return parse_driving_licence(text)


def parse_uk_passport(text: str) -> dict[str, str]:
    lines = [re.sub(r"[^A-Z0-9<]", "", line.upper()) for line in text.splitlines()]
    mrz_lines = [line for line in lines if "<" in line and len(line) >= 30]
    extracted = {"full_name": "", "id_number": "", "date_of_birth": ""}
    if mrz_lines:
        first_line = next((line for line in mrz_lines if line.startswith("P<")), mrz_lines[0])
        second_line = mrz_lines[1] if len(mrz_lines) > 1 else ""
        name_blob = first_line.split("GBR", 1)[-1]
        name_parts = [part for part in name_blob.replace("<<", "|").split("|") if part]
        extracted["full_name"] = " ".join(part.replace("<", " ").strip() for part in name_parts).title()
        if len(second_line) >= 19:
            extracted["id_number"] = second_line[:9].replace("<", "").strip()
            extracted["date_of_birth"] = parse_mrz_birth_date(second_line[13:19])
    return backfill_from_generic_patterns(text, extracted)


def parse_brp(text: str) -> dict[str, str]:
    extracted = {
        "full_name": extract_label_value(text, "Surname"),
        "id_number": extract_brp_number(text),
        "date_of_birth": extract_date(text),
    }
    given_names = extract_label_value(text, "Given Names")
    if given_names:
        extracted["full_name"] = " ".join(part for part in [given_names, extracted["full_name"]] if part).strip()
    return backfill_from_generic_patterns(text, extracted)


def parse_evisa(text: str) -> dict[str, str]:
    extracted = {
        "full_name": extract_label_value(text, "Name"),
        "id_number": extract_generic_id_number(text),
        "date_of_birth": extract_date(text),
    }
    return backfill_from_generic_patterns(text, extracted)


def parse_driving_licence(text: str) -> dict[str, str]:
    extracted = {
        "full_name": extract_driving_licence_name(text),
        "id_number": extract_driving_licence_number(text),
        "date_of_birth": extract_date(text),
    }
    return backfill_from_generic_patterns(text, extracted)


def extract_label_value(text: str, label: str) -> str:
    pattern = re.compile(rf"{re.escape(label)}[:\\s]+([A-Z][A-Z '\\-]+)", re.IGNORECASE)
    match = pattern.search(text)
    if not match:
        return ""
    return compact_spaces(match.group(1)).title()


def extract_brp_number(text: str) -> str:
    match = re.search(r"\b[A-Z]{2}\d{7}\b", text.upper())
    return match.group(0) if match else ""


def extract_driving_licence_name(text: str) -> str:
    surname = ""
    given_names = ""
    surname_match = re.search(r"1[.\s]+([A-Z][A-Z '\\-]+)", text.upper())
    if surname_match:
        surname = compact_spaces(surname_match.group(1)).title()
    given_match = re.search(r"2[.\s]+([A-Z][A-Z '\\-]+)", text.upper())
    if given_match:
        given_names = compact_spaces(given_match.group(1)).title()
    return " ".join(part for part in [given_names, surname] if part).strip()


def extract_driving_licence_number(text: str) -> str:
    match = re.search(r"\b[A-Z]{5}\d{6}[A-Z0-9]{2,5}\b", text.upper())
    return match.group(0) if match else ""


def backfill_from_generic_patterns(text: str, extracted: dict[str, str]) -> dict[str, str]:
    if not extracted.get("date_of_birth"):
        extracted["date_of_birth"] = extract_date(text)
    if not extracted.get("full_name"):
        extracted["full_name"] = extract_fallback_name(text)
    if not extracted.get("id_number"):
        extracted["id_number"] = extract_generic_id_number(text)
    return {key: compact_spaces(value) for key, value in extracted.items()}


def extract_date(text: str) -> str:
    date_patterns = [
        r"\b(\d{2}/\d{2}/\d{4})\b",
        r"\b(\d{2}-\d{2}-\d{4})\b",
        r"\b(\d{2}\s+[A-Z]{3}\s+\d{4})\b",
    ]
    for pattern in date_patterns:
        match = re.search(pattern, text.upper())
        if match:
            return normalise_detected_date(match.group(1))
    return ""


def normalise_detected_date(raw_value: str) -> str:
    candidates = ("%d/%m/%Y", "%d-%m-%Y", "%d %b %Y")
    for fmt in candidates:
        try:
            return datetime.strptime(raw_value, fmt).strftime("%d/%m/%Y")
        except ValueError:
            continue
    return raw_value


def parse_mrz_birth_date(raw_value: str) -> str:
    if len(raw_value) != 6 or not raw_value.isdigit():
        return ""
    year = int(raw_value[:2])
    month = int(raw_value[2:4])
    day = int(raw_value[4:6])
    current_year = date.today().year % 100
    century = 1900 if year > current_year else 2000
    try:
        parsed = date(century + year, month, day)
    except ValueError:
        return ""
    return parsed.strftime("%d/%m/%Y")


def extract_fallback_name(text: str) -> str:
    for line in text.splitlines():
        candidate = compact_spaces(line)
        if not candidate:
            continue
        if re.fullmatch(r"[A-Z][A-Z '\\-]{4,}", candidate.upper()):
            return candidate.title()
    return ""


def extract_generic_id_number(text: str) -> str:
    patterns = [
        r"\b\d{9}\b",
        r"\b[A-Z]{2}\d{7}\b",
        r"\b[A-Z]{5}\d{6}[A-Z0-9]{2,5}\b",
    ]
    upper_text = text.upper()
    for pattern in patterns:
        match = re.search(pattern, upper_text)
        if match:
            return match.group(0)
    return ""


def compact_spaces(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_root = Path(temp_dir)
        docx_path = temp_root / "agreement.docx"
        txt_path = temp_root / "agreement.txt"
        docx_path.write_bytes(docx_bytes)

        conversion = subprocess.run(
            [
                "textutil",
                "-convert",
                "txt",
                "-output",
                str(txt_path),
                str(docx_path),
            ],
            capture_output=True,
            text=True,
            check=False,
        )
        if conversion.returncode != 0:
            raise RuntimeError(conversion.stderr.strip() or "Unable to convert the document to text.")

        pdf_result = subprocess.run(
            ["/usr/sbin/cupsfilter", "-m", "application/pdf", str(txt_path)],
            capture_output=True,
            check=False,
        )
        if pdf_result.returncode != 0:
            message = pdf_result.stderr.decode("utf-8", errors="ignore").strip()
            raise RuntimeError(message or "Unable to convert the agreement to PDF.")
        return pdf_result.stdout
